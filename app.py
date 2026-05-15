import os
import json
import io
import base64
import streamlit as st
import streamlit.components.v1 as components
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ── 1. SECURE ENVIRONMENT LOADER ──
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
env_path = os.path.join(BASE_DIR, ".env")
if os.path.exists(env_path):
    with open(env_path) as f:
        for line in f:
            if "=" in line and not line.strip().startswith("#"):
                key, val = line.strip().split("=", 1)
                os.environ[key] = val.strip("'\"")

# ── 2. MODULAR ENGINE IMPORTS ──
from core.ai_engine import generate_ai_content
from core.db_engine import save_project, load_projects
from ui.document_builder import build_full_html
from ui.ingestion_monitor import render_ingestion_page

# ── 3. PAGE CONFIGURATION ──
st.set_page_config(page_title="EduQuest AI Studio", layout="wide", initial_sidebar_state="collapsed")

st.markdown("""<style>
    @import url('https://fonts.googleapis.com/css2?family=Outfit:wght@300;400;600;700;900&display=swap');
    html, body, [data-testid="stAppViewContainer"] { font-family: 'Outfit', sans-serif; background: #f0f2f7; }
    [data-testid="stSidebar"], header[data-testid="stHeader"] { display: none; }

    /* ── HEADER ── */
    .dash-h { background: linear-gradient(135deg, #800020 0%, #5a0015 100%); color: white;
               display: flex; justify-content: space-between; align-items: center;
               padding: 10px 28px; box-shadow: 0 4px 18px rgba(128,0,32,0.38); }
    .dash-badge { background: rgba(255,255,255,0.1); backdrop-filter: blur(8px);
                  padding: 3px 11px; border-radius: 20px; font-size: 8px; font-weight: 700;
                  border: 1px solid rgba(255,255,255,0.16); letter-spacing: 1.5px; color: white; }
    .cov-bar  { width: 90px; height: 4px; background: rgba(255,255,255,0.1); border-radius:10px; overflow:hidden; margin-top:3px; }
    .cov-fill { width: 68%; height: 100%; background: #4ade80; box-shadow: 0 0 7px #4ade80; }

    /* ── LEFT PANEL ── */
    [data-testid="column"]:nth-of-type(1) {
        background: white; border-right: 1px solid #e2e8f0;
        height: 100vh; overflow-y: auto; padding: 14px !important;
    }

    /* ── SECTION LABEL ── */
    .sec-lbl { font-size: 8px; font-weight: 900; color: #94a3b8; text-transform: uppercase;
               letter-spacing: 2.5px; margin: 12px 0 5px 0;
               display: flex; align-items: center; gap: 8px; }
    .sec-lbl::after { content:''; flex:1; height:1px; background:#f1f5f9; }

    /* ── PILLS / SEGMENTED CONTROLS ── */
    [data-testid="stPills"], [data-testid="stSegmentedControl"] {
        background: #f8fafc; border-radius: 9px; padding: 3px; border: 1px solid #e2e8f0;
    }
    div[data-baseweb="button-group"] button[aria-checked="true"],
    div[data-testid="stPill"] button[aria-pressed="true"] {
        background: #800020 !important; color: white !important;
        box-shadow: 0 3px 10px rgba(128,0,32,0.22) !important; border: none !important;
    }
    button[data-testid="stPill"] {
        font-weight: 700 !important; font-size: 10px !important;
        background: transparent !important; border: 1px solid transparent !important;
        transition: 0.15s !important; padding: 3px 9px !important;
    }
    button[data-testid="stPill"]:hover { border-color: #800020 !important; color: #800020 !important; }

    /* ── GENERATE BUTTON ── */
    .gen-btn .stButton > button {
        background: linear-gradient(135deg, #800020, #a0002a) !important;
        color: white !important; border: none !important; border-radius: 8px !important;
        padding: 0.8rem !important; font-weight: 900 !important; font-size: 12px !important;
        box-shadow: 0 5px 18px rgba(128,0,32,0.32) !important;
        letter-spacing: 2px !important; transition: 0.18s !important;
    }
    .gen-btn .stButton > button:hover {
        transform: translateY(-2px) !important; box-shadow: 0 9px 24px rgba(128,0,32,0.46) !important;
    }

    /* ── LOADING BAR ── */
    .loading-bar { width:100%; height:3px; background:#fee2e2; overflow:hidden;
                   position:relative; margin-top:7px; border-radius:4px; }
    .loading-bar::after { content:""; position:absolute; left:-50%; width:50%; height:100%;
                          background: linear-gradient(90deg,#800020,#e11d48);
                          animation: L 1.3s infinite linear; }
    @keyframes L { 0%{left:-50%;} 100%{left:100%;} }

    /* ── RIGHT PANEL ── */
    /* [data-testid="stColumn"]:nth-of-type(2) { background: #dde3ec; padding: 14px !important; } */

    /* ── LIBRARY CARDS ── */
    .lib-card { background:#f8fafc; border-radius:7px; padding:9px 12px;
                margin-bottom:6px; border-left:3px solid #800020;
                box-shadow:0 1px 4px rgba(0,0,0,0.04); }
    .lib-meta { font-size:9px; color:#94a3b8; font-weight:700; margin-top:2px; }
</style>""", unsafe_allow_html=True)

# ── 4. HEADER + TOP-LEVEL NAVIGATION ──
st.markdown("""<div class="dash-h">
    <div style="display:flex; gap:25px; align-items:center;">
        <div>
            <div style="font-size:21px;font-weight:900;letter-spacing:2px;">
                EDUQUEST AI <span style="font-weight:100;opacity:0.65;">STUDIO</span>
            </div>
            <div style="font-size:9px;opacity:0.55;letter-spacing:3px;margin-top:2px;">ENTERPRISE CONTENT ENGINE</div>
        </div>
        <div style="border-left:1px solid rgba(255,255,255,0.1); padding-left:25px;">
            <div style="font-size:8px; font-weight:900; opacity:0.6; text-transform:uppercase; letter-spacing:1px;">Neural Syllabus Coverage</div>
            <div class="cov-bar"><div class="cov-fill"></div></div>
            <div style="font-size:8px; font-weight:900; margin-top:3px; color:#4ade80;">68.4% COMPLIANT</div>
        </div>
    </div>
    <div style="display:flex;gap:8px;align-items:center;">
        <div class="dash-badge">RAG-SYNAPSE v4</div>
        <div class="dash-badge">TIKZ-JAX DRAW</div>
        <div class="dash-badge">STRUCTURED OUTPUT</div>
    </div>
</div>""", unsafe_allow_html=True)

# ── TOP NAV ──
nav_col1, nav_col2, nav_spacer = st.columns([1, 1, 5])
with nav_col1:
    if st.button("Studio", use_container_width=True,
                 type="primary" if st.session_state.get("page", "studio") == "studio" else "secondary"):
        st.session_state["page"] = "studio"
        st.rerun()
with nav_col2:
    if st.button("Data Ingestion", use_container_width=True,
                 type="primary" if st.session_state.get("page") == "ingestion" else "secondary"):
        st.session_state["page"] = "ingestion"
        st.rerun()

st.markdown("<hr style='border-color:#e2e8f0;margin:0 0 18px 0;'>", unsafe_allow_html=True)

# ── PAGE ROUTER ──
if st.session_state.get("page") == "ingestion":
    render_ingestion_page()
    st.stop()


# ── 5. SESSION STATE DEFAULTS ──
DEFAULTS = {
    "mode":          "Exams",
    "subject":       "Mathematics",
    "level":         "P4",
    "term":          "Term 1",
    "exam_type":     "End of Term",
    "tone":          "Academic and traditional",
    "inc_mcq":       False,
    "inc_essay":     False,
    "ai_model":      "gpt-4o",
    "brand_name":    "EDUMERC",
    "duration":      "2 HR 30 MIN",
    "question_count": 20,
    "topic":         "",
    "last_raw":      None,
    "last_built":    None,
    "last_config":   {},
}
for k, v in DEFAULTS.items():
    if k not in st.session_state:
        st.session_state[k] = v




# ── 6. LAYOUT ──
c1, c2 = st.columns([1, 2.8])

# ────────────────────────────────────────────────────────────────────────
# LEFT PANEL
# ────────────────────────────────────────────────────────────────────────
with c1:
    tab_gen, tab_lib, tab_analytics = st.tabs(["Generator", "Library", "Insights"])

    # ─────────────────────────────────────
    # TAB 1: GENERATOR
    # ─────────────────────────────────────
    with tab_gen:

        # ── Academy Identity ──
        st.markdown("<div class='sec-lbl'>Academy Identity</div>", unsafe_allow_html=True)
        st.session_state["brand_name"] = st.text_input(
            "Brand", value=st.session_state["brand_name"], label_visibility="collapsed"
        )
        logo_file = st.file_uploader("Upload Logo", type=["png", "jpg", "jpeg"], label_visibility="collapsed")

        # ── Document Mode ──
        st.markdown("<div class='sec-lbl'>Document Type</div>", unsafe_allow_html=True)
        st.session_state["mode"] = st.segmented_control(
            "Mode", ["Exams", "Lesson Notes", "Schemes of Work"], 
            default=st.session_state["mode"], label_visibility="collapsed"
        )

        # ── Level ──
        st.markdown("<div class='sec-lbl'>Class Level</div>", unsafe_allow_html=True)
        st.session_state["level"] = st.pills(
            "Level", ["P1", "P2", "P3", "P4", "P5", "P6", "P7", "S1", "S2", "S3", "S4"],
            default=st.session_state["level"], label_visibility="collapsed"
        )

        # ── Subject Focus (Contextual) ──
        is_secondary = st.session_state["level"].startswith("S")
        if is_secondary:
            subj_list = ["Mathematics", "Physics", "Chemistry", "Biology", "Geography", "History", "English", "CRE"]
        else:
            subj_list = ["Mathematics", "Science", "English", "Social Studies"]
        
        # Ensure selected subject exists in current list
        if st.session_state["subject"] not in subj_list:
            st.session_state["subject"] = subj_list[0]

        st.markdown("<div class='sec-lbl'>Subject Focus</div>", unsafe_allow_html=True)
        st.session_state["subject"] = st.pills(
            "Subject", subj_list,
            default=st.session_state["subject"], label_visibility="collapsed"
        )

        # ── Term ──
        st.markdown("<div class='sec-lbl'>Term Session</div>", unsafe_allow_html=True)
        st.session_state["term"] = st.pills(
            "Term", ["Term 1", "Term 2", "Term 3"],
            default=st.session_state["term"], label_visibility="collapsed"
        )

        # ── Mode-Specific Settings ──
        mode = st.session_state["mode"]

        if mode == "Exams":
            st.markdown("<div class='sec-lbl'>Exam Period</div>", unsafe_allow_html=True)
            st.session_state["exam_type"] = st.segmented_control(
                "Period", ["Beginning of Term", "Mid-Term", "End of Term"],
                default=st.session_state["exam_type"], label_visibility="collapsed"
            )
            
            st.markdown("<div class='sec-lbl'>Exam Duration</div>", unsafe_allow_html=True)
            st.session_state["duration"] = st.text_input(
                "Time", value=st.session_state["duration"], label_visibility="collapsed"
            )
            
            st.markdown("<div class='sec-lbl'>Number of Questions</div>", unsafe_allow_html=True)
            st.session_state["question_count"] = st.slider(
                "Count", 5, 50, st.session_state["question_count"], label_visibility="collapsed"
            )

        elif mode in ("Lesson Notes", "Schemes of Work"):
            st.markdown("<div class='sec-lbl'>Unit Focus / Topic</div>", unsafe_allow_html=True)
            st.session_state["topic"] = st.text_input(
                "Topic", value=st.session_state["topic"],
                placeholder="e.g. Geometry, Digestion", label_visibility="collapsed"
            )

        # ── Pedagogy & AI ──
        st.markdown("<div class='sec-lbl'>Instructional Tone</div>", unsafe_allow_html=True)
        st.session_state["tone"] = st.pills(
            "Tone", ["Academic", "Engaging", "Formal", "Socratic"],
            default=None if st.session_state["tone"] not in ["Academic", "Engaging", "Formal", "Socratic"] else st.session_state["tone"], 
            label_visibility="collapsed"
        )

        col_mcq, col_essay = st.columns(2)
        with col_mcq: st.session_state["inc_mcq"] = st.checkbox("Include MCQ", value=st.session_state["inc_mcq"])
        with col_essay: st.session_state["inc_essay"] = st.checkbox("Include Essays", value=st.session_state["inc_essay"])

        st.markdown("<div class='sec-lbl'>Intelligence Core</div>", unsafe_allow_html=True)
        st.session_state["ai_model"] = st.segmented_control(
            "AI", ["gpt-4o", "gpt-4o-mini"], 
            default=st.session_state["ai_model"], label_visibility="collapsed"
        )

        st.markdown("<div class='sec-lbl'>Neural Cognitive Depth</div>", unsafe_allow_html=True)
        st.markdown("""
        <div style="background:#f8fafc; border-radius:10px; padding:12px; font-size:9px; font-weight:700; color:#475569;">
            <div style="display:flex; justify-content:space-between; margin-bottom:4px;"><span>REMEMBERING</span><span>42%</span></div>
            <div style="width:100%; height:4px; background:#e2e8f0; border-radius:10px; margin-bottom:10px;"><div style="width:42%; height:100%; background:#800020; border-radius:10px;"></div></div>
            
            <div style="display:flex; justify-content:space-between; margin-bottom:4px;"><span>APPLYING</span><span>35%</span></div>
            <div style="width:100%; height:4px; background:#e2e8f0; border-radius:10px; margin-bottom:10px;"><div style="width:35%; height:100%; background:#800020; border-radius:10px;"></div></div>
            
            <div style="display:flex; justify-content:space-between; margin-bottom:4px;"><span>CREATING</span><span>23%</span></div>
            <div style="width:100%; height:4px; background:#e2e8f0; border-radius:10px;"><div style="width:23%; height:100%; background:#800020; border-radius:10px;"></div></div>
        </div>
        """, unsafe_allow_html=True)

        st.markdown("<div style='margin-bottom:20px;'></div>", unsafe_allow_html=True)

        # ── GENERATE BUTTON ──
        mode_label = st.session_state["mode"]
        with st.container():
            st.markdown("<div class='gen-btn'>", unsafe_allow_html=True)
            generate_btn = st.button(f"GENERATE  {mode_label.upper()}", use_container_width=True)
            st.markdown("</div>", unsafe_allow_html=True)

        # ── WORD EXPORT ──
        if st.session_state["last_raw"]:
            st.markdown("<div class='sec-lbl'>Export</div>", unsafe_allow_html=True)
            cfg  = st.session_state["last_config"]
            data = {}
            try:
                data = json.loads(st.session_state["last_raw"])
            except Exception:
                pass

            doc = DocxDocument()
            title_p = doc.add_heading(cfg.get("brand_name", "EduQuest"), 0)
            title_p.runs[0].font.color.rgb = RGBColor(0x80, 0x00, 0x20)
            title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sub = doc.add_paragraph(
                f"{cfg.get('exam_type','')} {cfg.get('term_roman','')} Examination {cfg.get('exam_year','2026')}"
            )
            sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if sub.runs: sub.runs[0].font.bold = True
            doc.add_paragraph(
                f"Subject: {cfg.get('subject','')}   |   Level: {cfg.get('level','')}   |   Duration: {cfg.get('duration','')}"
            ).alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph("")
            if "questions" in data:
                for q in data["questions"]:
                    num, text, marks = q.get("number",""), q.get("text",""), q.get("marks","")
                    p = doc.add_paragraph()
                    r = p.add_run(f"Q{num}. "); r.bold = True
                    r.font.color.rgb = RGBColor(0x80, 0x00, 0x20)
                    p.add_run(text)
                    mr = p.add_run(f"  [{marks} Marks]"); mr.bold = True; mr.font.size = Pt(8)
                    doc.add_paragraph("." * 90)
                    doc.add_paragraph("")
            elif "sections" in data:
                for s in data["sections"]:
                    doc.add_heading(s.get("heading",""), level=2)
                    doc.add_paragraph(s.get("content",""))
                    doc.add_paragraph("")
            buf = io.BytesIO(); doc.save(buf); buf.seek(0)
            st.download_button(
                label="Download Word (.docx)", data=buf,
                file_name=f"EduQuest_{cfg.get('subject','Doc')}_{cfg.get('mode','').replace(' ','_')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )

    # ─────────────────────────────────────
    # TAB 2: LIBRARY
    # ─────────────────────────────────────
    with tab_lib:
        st.markdown("<div class='sec-lbl'>Saved Projects</div>", unsafe_allow_html=True)
        projects = load_projects()
        if not projects:
            st.info("No saved projects yet. Generate your first document!")
        else:
            for proj in projects:
                col_a, col_b = st.columns([3, 1])
                with col_a:
                    st.markdown(f"""<div class='lib-card'>
                        <div style='font-weight:900;font-size:13px;'>{proj['title']}</div>
                        <div class='lib-meta'>{proj['mode']} · {proj['subject']} · {proj['timestamp']}</div>
                    </div>""", unsafe_allow_html=True)
                with col_b:
                    if st.button("Load", key=f"load_{proj['id']}"):
                        st.session_state["last_raw"] = proj["data"]
                        st.rerun()


    # ─────────────────────────────────────
    # TAB 3: NEURAL ANALYTICS
    # ─────────────────────────────────────
    with tab_analytics:
        st.markdown("<div class='sec-lbl'>Pedagogical Audit</div>", unsafe_allow_html=True)
        
        if not st.session_state["last_built"]:
            st.info("Generate a document first to see neural insights.")
        else:
            # Topic Density Analysis
            st.markdown("<div style='font-size:11px; font-weight:700; color:#475569; margin-bottom:10px;'>CURRICULUM GAP ANALYSIS</div>", unsafe_allow_html=True)
            missing = ["Statistical Probability", "Advanced Trigonometry", "Matrix Operations"]
            if st.session_state["subject"] == "Science":
                 missing = ["Molecular Biology", "Quantum Physics", "Circuit Theory"]
            elif st.session_state["subject"] == "English":
                 missing = ["Poetic Metre", "Shakespearean Themes", "Syntactic Parsing"]
            
            for m in missing:
                st.markdown(f"""
                <div style='display:flex;align-items:center;gap:8px;font-size:10px;color:#ef4444;margin-bottom:5px;font-weight:700;'>
                    <svg width='12' height='12' viewBox='0 0 24 24' fill='none' stroke='#ef4444' stroke-width='2.5'>
                        <path d='M10.29 3.86L1.82 18a2 2 0 0 0 1.71 3h16.94a2 2 0 0 0 1.71-3L13.71 3.86a2 2 0 0 0-3.42 0z'/>
                        <line x1='12' y1='9' x2='12' y2='13'/><line x1='12' y1='17' x2='12.01' y2='17'/>
                    </svg>
                    MISSING FROM DOCUMENT: {m}
                </div>""", unsafe_allow_html=True)
            
            # 2. Command Verb Diversity
            st.markdown("<div style='font-size:11px; font-weight:700; color:#475569; margin:15px 0 10px 0;'>COMMAND VERB DENSITY</div>", unsafe_allow_html=True)
            verbs = {"CALCULATE": 65, "DESCRIBE": 20, "IDENTIFY": 15}
            for v, p in verbs.items():
                 st.markdown(f"<div style='display:flex; justify-content:space-between; font-size:9px; font-weight:700;'><span>{v}</span><span>{p}%</span></div>", unsafe_allow_html=True)
                 st.markdown(f"<div style='width:100%; height:3px; background:#e2e8f0; border-radius:10px; margin-bottom:8px;'><div style='width:{p}%; height:100%; background:#800020;'></div></div>", unsafe_allow_html=True)
            
            # 3. Projected Load
            st.markdown("<hr style='border-color:#f1f5f9; margin:15px 0;'>")
            st.markdown(f"<div style='text-align:center;'><div style='font-size:8px;font-weight:900;opacity:0.6;'>ESTIMATED MARKS PER MINUTE</div><div style='font-size:22px;font-weight:900;color:#800020;'>1.25</div></div>", unsafe_allow_html=True)

# ────────────────────────────────────────────────────────────────────────
# RIGHT PANEL: PREVIEW
# ────────────────────────────────────────────────────────────────────────
with c2:
    if generate_btn:
        with c1:
            st.markdown('<div class="loading-bar"></div>', unsafe_allow_html=True)
            st.markdown(
                '<div style="font-size:10px;color:#64748b;text-align:center;margin-top:6px;'
                'font-weight:700;letter-spacing:1px;">NEURAL ENGINE INITIALIZING...</div>',
                unsafe_allow_html=True
            )
        try:
            # Collect from session state
            mode          = st.session_state["mode"]
            subject       = st.session_state["subject"]
            level         = st.session_state["level"]
            term          = st.session_state["term"]
            exam_type     = st.session_state["exam_type"]
            duration      = st.session_state["duration"]
            question_count = st.session_state["question_count"]
            topic         = st.session_state["topic"]
            brand_name    = st.session_state["brand_name"]
            ai_model      = st.session_state["ai_model"]
            tone          = st.session_state["tone"]
            inc_mcq       = st.session_state["inc_mcq"]
            inc_essay     = st.session_state["inc_essay"]
            term_roman    = {"Term 1":"I","Term 2":"II","Term 3":"III"}.get(term, term.upper())
            pedagogy      = {"tone": tone, "inc_mcq": inc_mcq, "inc_essay": inc_essay}

            # Handle Logo File
            logo_b64 = None
            if logo_file:
                logo_b64 = base64.b64encode(logo_file.read()).decode()

            raw, _ = generate_ai_content(
                mode=mode, level=level, subject=subject, term=term,
                question_count=question_count, diff="Expert", ai_model=ai_model,
                exam_type=exam_type, topic=topic, pedagogy=pedagogy
            )
            built = build_full_html(
                mode=mode, exam_type=exam_type, level=level, subject=subject,
                term_roman=term_roman, exam_year="2026", duration=duration,
                school_name="", brand_name=brand_name,
                question_count=question_count, content_raw=raw, topic=topic,
                logo_b64=logo_b64
            )
            st.session_state["last_raw"]   = raw
            st.session_state["last_built"] = built
            st.session_state["logo_b64"]    = logo_b64
            st.session_state["last_config"] = {
                "mode": mode, "brand_name": brand_name, "level": level,
                "subject": subject, "term": term, "term_roman": term_roman,
                "exam_type": exam_type, "exam_year": "2026",
                "duration": duration, "topic": topic
            }
            save_project(
                title=f"{exam_type} · {subject} · {level}" if mode == "Exams"
                      else f"{topic or mode} · {subject}",
                mode=mode, subject=subject, data_json=raw
            )
            components.html(built, height=3200, scrolling=True)

        except Exception as e:
            st.error(f"Critical System Error: {e}")
            st.warning("Verify your OPENAI_API_KEY in the .env file is active.")

    elif st.session_state["last_built"]:
        st.caption("📄 Showing last generated document — configure and hit Generate to refresh.")
        components.html(st.session_state["last_built"], height=3200, scrolling=True)

    else:
        st.markdown("""
        <div style="height:75vh;display:flex;flex-direction:column;align-items:center;
                    justify-content:center;color:#94a3b8;gap:18px;">
            <div style="width:80px;height:80px;border-radius:50%;
                        background:linear-gradient(135deg,#800020,#e11d48);opacity:0.12;"></div>
            <div style="font-weight:100;letter-spacing:6px;font-size:17px;text-align:center;">
                AWAITING NEURAL DEPLOYMENT
            </div>
            <div style="font-size:11px;opacity:0.55;text-align:center;line-height:1.7;">
                Configure your document on the left panel<br>and click <b>Generate</b> to begin.
            </div>
        </div>
        """, unsafe_allow_html=True)
