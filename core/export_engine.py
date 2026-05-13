import io
import json
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_docx_stream(data_json: str, config: dict):
    """
    Generates a Microsoft Word (.docx) document from EduQuest JSON data.
    """
    try:
        data = json.loads(data_json)
    except Exception as e:
        raise ValueError(f"Invalid JSON data for export: {e}")

    doc = DocxDocument()
    
    # ── GLOBAL STYLING ──
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    def add_markdown(paragraph, text):
        """Parses basic markdown **bold** and strips KaTeX $ symbols for cleaner Word output."""
        clean_text = text.replace("$", "")
        parts = clean_text.split("**")
        for i, part in enumerate(parts):
            if part:
                run = paragraph.add_run(part)
                if i % 2 == 1:
                    run.bold = True

    # ── HEADER ──
    brand_name = config.get("brand_name", "EDUMERC")
    title_p = doc.add_heading(brand_name, 0)
    if title_p.runs:
        title_p.runs[0].font.color.rgb = RGBColor(0x80, 0x00, 0x20) # Burgundy
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub_title = f"{config.get('exam_type', 'Examination')} {config.get('term_roman', '')} {config.get('exam_year', '2026')}"
    sub = doc.add_paragraph(sub_title)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if sub.runs:
        sub.runs[0].font.bold = True
        sub.runs[0].font.size = Pt(14)

    meta = f"Subject: {config.get('subject', 'N/A')}   |   Level: {config.get('level', 'N/A')}   |   Duration: {config.get('duration', '2 HR 30 MIN')}"
    meta_p = doc.add_paragraph(meta)
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph("_" * 60).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")
    
    # ── CANDIDATE DETAILS BOX ──
    cand_table = doc.add_table(rows=2, cols=2)
    cand_table.style = 'Table Grid'
    
    r0 = cand_table.rows[0].cells
    r0[0].text = " CANDIDATE NAME:"
    r0[0].paragraphs[0].runs[0].font.bold = True
    r0[1].text = "" # Space for student to write
    
    r1 = cand_table.rows[1].cells
    r1[0].text = " INDEX NUMBER:"
    r1[0].paragraphs[0].runs[0].font.bold = True
    r1[1].text = " SIGNATURE:"
    r1[1].paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph("")
    doc.add_paragraph("_" * 60).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")

    # ── CONTENT ──
    if "scenario_text" in data:
        doc.add_heading("COMPETENCY SCENARIO", level=2)
        p = doc.add_paragraph(data["scenario_text"])
        p.italic = True
        doc.add_paragraph("")

    if "questions" in data:
        for q in data["questions"]:
            num = q.get("number", "")
            text = q.get("text", "")
            marks = q.get("marks", "")
            
            p = doc.add_paragraph()
            r = p.add_run(f"Q{num}. ")
            r.bold = True
            r.font.color.rgb = RGBColor(0x80, 0x00, 0x20)
            
            p.add_run("  ")
            add_markdown(p, text)
            
            mr = p.add_run(f"  [{marks} Marks]")
            mr.bold = True
            mr.font.size = Pt(9)
            
            # Embed generated images if available
            tikz = q.get("tikz_code", "")
            if tikz and "<img" in tikz:
                import re, os
                match = re.search(r'src="([^"]+)"', tikz)
                if match:
                    src = match.group(1)
                    if src.startswith("/generated/"):
                        img_path = os.path.join("frontend", "public", src.lstrip("/"))
                        if os.path.exists(img_path):
                            pic_p = doc.add_paragraph()
                            pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            pic_r = pic_p.add_run()
                            pic_r.add_picture(img_path, width=Inches(3.5))

            if "options" in q and q["options"]:
                for idx, opt in enumerate(q["options"]):
                    letter = chr(65 + idx)
                    opt_p = doc.add_paragraph(f"    {letter}. ")
                    add_markdown(opt_p, opt)

            # Answer lines
            doc.add_paragraph("")
            doc.add_paragraph("." * 100)
            doc.add_paragraph("")
    
    elif "sections" in data:
        for s in data["sections"]:
            doc.add_heading(s.get("heading", "Section"), level=2)
            doc.add_paragraph(s.get("content", ""))
            doc.add_paragraph("")

    # ── TEACHER'S MARKING GUIDE (Optional/Confidential) ──
    if "questions" in data:
        doc.add_page_break()
        guide_h = doc.add_heading("TEACHER'S MARKING GUIDE (CONFIDENTIAL)", level=1)
        guide_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'QN'
        hdr_cells[1].text = 'EXPECTED ANSWER / WORKING'
        hdr_cells[2].text = 'MARKS'
        
        for q in data["questions"]:
            row_cells = table.add_row().cells
            row_cells[0].text = str(q.get("number", ""))
            row_cells[1].text = q.get("answer", "N/A")
            row_cells[2].text = str(q.get("marks", ""))

    # ── STREAM ──
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
